#!/usr/bin/env python3
"""
Convert Markdown to DOCX using python-docx and markdown libraries.
"""

import os
import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(document):
    """Add a page break to the document."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
            bottom={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
            start={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
            end={"sz": 12, "val": "single", "color": "#000000", "space": "0"}
        )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for side, attrs in kwargs.items():
        tag = {
            'top': 'w:top',
            'start': 'w:start',
            'left': 'w:left',
            'bottom': 'w:bottom',
            'end': 'w:end',
            'right': 'w:right',
            'insideH': 'w:insideH',
            'insideV': 'w:insideV',
        }.get(side)
        
        if tag:
            element = tcPr.find(tag)
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            
            for key, value in attrs.items():
                element.set(qn('w:{}'.format(key)), value)

def convert_markdown_to_docx(md_file, docx_file):
    """Convert Markdown file to DOCX."""
    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Split the content by page breaks
    pages = md_content.split('<div style="page-break-after: always;"></div>')
    
    # Create a new Document
    doc = Document()
    
    # Set document properties
    core_properties = doc.core_properties
    core_properties.title = "RRDM Release Notes Requirements"
    core_properties.author = "Cascade AI"
    
    # Process each page
    for i, page_content in enumerate(pages):
        # For all pages except the first, add a page break
        if i > 0:
            doc.add_page_break()
        
        # Parse markdown content
        lines = page_content.strip().split('\n')
        
        # Process lines
        j = 0
        while j < len(lines):
            line = lines[j].strip()
            
            # Handle headings
            if line.startswith('# '):
                # Title (H1)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line[2:])
                run.bold = True
                run.font.size = Pt(18)
                j += 1
            elif line.startswith('## '):
                # Section heading (H2)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(line[3:])
                run.bold = True
                run.font.size = Pt(16)
                j += 1
            elif line.startswith('### '):
                # Subsection heading (H3)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(line[4:])
                run.bold = True
                run.font.size = Pt(14)
                j += 1
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
                j += 1
            elif line.startswith('- '):
                # Bullet points
                p = doc.add_paragraph(line[2:], style='List Bullet')
                j += 1
            elif line.startswith('  * '):
                # Sub-bullet points
                p = doc.add_paragraph(line[4:], style='List Bullet 2')
                j += 1
            elif line.startswith('---'):
                # Horizontal rule
                doc.add_paragraph('_' * 50)
                j += 1
            elif line == '':
                # Empty line
                doc.add_paragraph()
                j += 1
            else:
                # Regular paragraph
                p = doc.add_paragraph(line)
                j += 1
    
    # Save the document
    doc.save(docx_file)
    print(f"Converted {md_file} to {docx_file}")

if __name__ == "__main__":
    import sys
    import docx
    
    md_file = "/Users/freddieo/Documents/rrdm/structured_requirements.md"
    docx_file = "/Users/freddieo/Documents/rrdm/structured_requirements.docx"
    
    convert_markdown_to_docx(md_file, docx_file)
