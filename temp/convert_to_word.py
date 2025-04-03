import re
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

# Create a new Document
doc = Document()

# Set up document properties
doc.core_properties.title = "Release Notes Requirements"
doc.core_properties.author = "Cascade AI"
doc.core_properties.created = datetime.datetime(2025, 3, 30)

# Set up styles
styles = doc.styles

# Title style
title_style = styles['Title']
title_style.font.name = 'Calibri'
title_style.font.size = Pt(24)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0, 0, 0)

# Heading 1 style
heading1_style = styles['Heading 1']
heading1_style.font.name = 'Calibri'
heading1_style.font.size = Pt(18)
heading1_style.font.bold = True
heading1_style.font.color.rgb = RGBColor(0, 51, 102)

# Heading 2 style
heading2_style = styles['Heading 2']
heading2_style.font.name = 'Calibri'
heading2_style.font.size = Pt(16)
heading2_style.font.bold = True
heading2_style.font.color.rgb = RGBColor(0, 102, 153)

# Heading 3 style
heading3_style = styles['Heading 3']
heading3_style.font.name = 'Calibri'
heading3_style.font.size = Pt(14)
heading3_style.font.bold = True
heading3_style.font.color.rgb = RGBColor(0, 102, 204)

# Normal text style
normal_style = styles['Normal']
normal_style.font.name = 'Calibri'
normal_style.font.size = Pt(11)
normal_style.font.color.rgb = RGBColor(0, 0, 0)

# Read the text file
with open('/Users/freddieo/Documents/rrdm/release_notes_requirements.txt', 'r') as file:
    content = file.read()

# Create a cover page
doc.add_paragraph("Release Notes Requirements", 'Title')
doc.add_paragraph("Document Version: 2.0", 'Subtitle')
doc.add_paragraph("Date: March 30, 2025", 'Subtitle')
doc.add_paragraph("Author: Cascade AI", 'Subtitle')

# Add a page break after the cover page
doc.add_page_break()

# Create a table of contents page
doc.add_paragraph("Table of Contents", 'Title')
doc.add_paragraph("This document contains the following sections:", 'Normal')

# Extract main sections for TOC
sections = re.findall(r'## \d+\. (.*?)(?=\n)', content)
for i, section in enumerate(sections):
    p = doc.add_paragraph(style='Normal')
    p.add_run(f"{i+1}. {section}")

# Add a page break after the TOC
doc.add_page_break()

# Split content by main sections (level 2 headings)
main_sections = re.split(r'(?=## \d+\.)', content)

# Skip the first element if it's just the document title and metadata
if not main_sections[0].strip().startswith('##'):
    main_sections = main_sections[1:]

# Process each main section
for section in main_sections:
    if not section.strip():
        continue
        
    # Extract section title
    section_title_match = re.match(r'## \d+\. (.*?)(?=\n)', section)
    if section_title_match:
        section_title = section_title_match.group(1)
        
        # Add the section title
        doc.add_paragraph(section_title, 'Heading 1')
        
        # Split the section into subsections
        subsections = re.split(r'(?=### )', section)
        
        # Skip the section title part
        subsection_content = subsections[0].split('\n', 1)[1] if len(subsections) > 0 else ""
        if subsection_content.strip():
            for line in subsection_content.strip().split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip(), 'Normal')
        
        # Process each subsection (starting from index 1 to skip the title part)
        for i in range(1, len(subsections)):
            subsection = subsections[i]
            
            # Extract subsection title
            subsection_title_match = re.match(r'### (.*?)(?=\n)', subsection)
            if subsection_title_match:
                subsection_title = subsection_title_match.group(1)
                
                # Add the subsection title
                doc.add_paragraph(subsection_title, 'Heading 2')
                
                # Get subsection content (skip the title)
                subsection_content = subsection.split('\n', 1)[1] if len(subsection.split('\n')) > 1 else ""
                
                # Process content with proper formatting
                current_list = None
                for line in subsection_content.strip().split('\n'):
                    if line.strip():
                        if line.strip().startswith('**') and line.strip().endswith(':**'):
                            # This is a bold heading (like "Implementation Requirements:")
                            heading_text = line.strip().strip('**:')
                            doc.add_paragraph(heading_text, 'Heading 3')
                        elif line.strip().startswith('- '):
                            # This is a list item
                            if current_list is None:
                                current_list = True
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(line.strip()[2:])
                        else:
                            # Regular paragraph
                            current_list = None
                            p = doc.add_paragraph(style='Normal')
                            
                            # Check for bold text within the line
                            bold_parts = re.findall(r'\*\*(.*?)\*\*', line)
                            if bold_parts:
                                # Line contains bold text
                                parts = re.split(r'(\*\*.*?\*\*)', line)
                                for part in parts:
                                    if part.startswith('**') and part.endswith('**'):
                                        # Bold part
                                        p.add_run(part.strip('**')).bold = True
                                    else:
                                        # Regular part
                                        p.add_run(part)
                            else:
                                # Regular line without formatting
                                p.add_run(line)
        
        # Add a page break after each main section
        doc.add_page_break()

# Save the document
doc.save('/Users/freddieo/Documents/rrdm/release_notes_requirements.docx')

print("Conversion complete. Document saved as 'release_notes_requirements.docx'")
